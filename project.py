from flask import Flask, render_template, request, send_file, redirect
from docx2pdf import convert
from reportlab.pdfgen import canvas
from docx import Document
import os
import PyPDF2
from pdf2docx import Converter
from io import StringIO

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def txt_to_pdf(input_path, output_path):
    c = canvas.Canvas(output_path)
    with open(input_path, 'r', encoding='utf-8') as file:
        text = file.read()
    c.drawString(50, 750, text)
    c.save()

def txt_to_docx(input_path, output_path):
    doc = Document()
    with open(input_path, 'r', encoding='utf-8') as file:
        doc.add_paragraph(file.read())
    doc.save(output_path)

def pdf_to_txt(input_path, output_path):
    with open(input_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = "\n".join([page.extract_text() for page in reader.pages])
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)

def pdf_to_docx(input_path, output_path):
    cv = Converter(input_path)
    cv.convert(output_path)
    cv.close()

def docx_to_txt(input_path, output_path):
    doc = Document(input_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)

@app.route('/', methods=['GET', 'POST'])
def home():
    if request.method == 'POST':
        if 'file' not in request.files:
            return redirect(request.url)
        
        file = request.files['file']
        convert_to = request.form.get('convert_to')
        
        if file.filename == '':
            return redirect(request.url)
        
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(input_path)
        
        try:
            output_filename = None
            output_path = None

            if convert_to == 'pdf':
                if file.filename.endswith('.docx'):
                    output_filename = file.filename.replace('.docx', '.pdf')
                    output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
                    convert(input_path, output_path)
                elif file.filename.endswith('.txt'):
                    output_filename = file.filename.replace('.txt', '.pdf')
                    output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
                    txt_to_pdf(input_path, output_path)
            
            elif convert_to == 'docx':
                if file.filename.endswith('.pdf'):
                    output_filename = file.filename.replace('.pdf', '.docx')
                    output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
                    pdf_to_docx(input_path, output_path)
                elif file.filename.endswith('.txt'):
                    output_filename = file.filename.replace('.txt', '.docx')
                    output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
                    txt_to_docx(input_path, output_path)
            
            elif convert_to == 'txt':
                if file.filename.endswith('.pdf'):
                    output_filename = file.filename.replace('.pdf', '.txt')
                    output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
                    pdf_to_txt(input_path, output_path)
                elif file.filename.endswith('.docx'):
                    output_filename = file.filename.replace('.docx', '.txt')
                    output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
                    docx_to_txt(input_path, output_path)
            
            os.remove(input_path)
            return send_file(output_path, as_attachment=True)
        
        except Exception as e:
            return f"Ошибка при конвертации: {e}"
    
    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True)